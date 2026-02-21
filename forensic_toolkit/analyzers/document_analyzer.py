"""
Document Analyzer
==================
Deep forensic analysis of document files.
Supports: PDF, DOCX, XLSX, PPTX, ODT, TXT, CSV

Uses: PyPDF2, python-docx, openpyxl
"""

import os
import zipfile
import json

from core.hasher import compute_hashes, compute_partial_hashes
from core.file_inspector import (
    get_file_info,
    detect_file_type,
    detect_appended_data,
    analyze_byte_distribution,
)


def analyze_document(file_path):
    """Perform full forensic analysis on a document file."""
    report = {
        "analysis_type": "Document Forensic Analysis",
        "file_info": {},
        "file_type_detection": {},
        "hashes": {},
        "partial_hashes": {},
        "document_properties": {},
        "metadata": {},
        "structure_analysis": {},
        "appended_data": [],
        "byte_distribution": {},
        "anomalies": [],
    }

    # Basic analysis
    report["file_info"] = get_file_info(file_path)
    report["file_type_detection"] = detect_file_type(file_path)
    report["hashes"] = compute_hashes(file_path)
    report["partial_hashes"] = compute_partial_hashes(file_path)
    report["byte_distribution"] = analyze_byte_distribution(file_path)
    report["appended_data"] = detect_appended_data(file_path)

    if report["appended_data"]:
        for item in report["appended_data"]:
            report["anomalies"].append(
                f"APPENDED DATA: {item['type']} - {item['extra_human']} extra"
            )

    ext = os.path.splitext(file_path)[1].lower()
    detected = report["file_type_detection"]["detected_type"]

    # Route to specific analyzer
    if ext == ".pdf" or "PDF" in detected:
        _analyze_pdf(file_path, report)
    elif ext == ".docx" or ("ZIP" in detected and ext == ".docx"):
        _analyze_docx(file_path, report)
    elif ext == ".xlsx" or ("ZIP" in detected and ext == ".xlsx"):
        _analyze_xlsx(file_path, report)
    elif ext in (".txt", ".csv", ".log", ".json", ".xml", ".html"):
        _analyze_text(file_path, report)

    # If it's a ZIP-based format, analyze the ZIP structure
    if "ZIP" in detected or ext in (".docx", ".xlsx", ".pptx", ".odt", ".ods"):
        _analyze_zip_structure(file_path, report)

    return report


def _analyze_pdf(file_path, report):
    """PDF-specific analysis."""
    try:
        from PyPDF2 import PdfReader

        reader = PdfReader(file_path)

        report["document_properties"] = {
            "page_count": len(reader.pages),
            "is_encrypted": reader.is_encrypted,
        }

        # Document metadata
        meta = reader.metadata
        if meta:
            metadata = {}
            for key in [
                "/Title", "/Author", "/Subject", "/Creator",
                "/Producer", "/CreationDate", "/ModDate",
                "/Keywords", "/Trapped",
            ]:
                val = meta.get(key)
                if val:
                    metadata[key.strip("/")] = str(val)
            report["metadata"] = metadata if metadata else {"note": "No metadata"}
        else:
            report["metadata"] = {"note": "No metadata available"}

        # Page analysis
        pages_info = []
        for i, page in enumerate(reader.pages[:10]):  # First 10 pages
            page_info = {
                "page_number": i + 1,
                "has_text": bool(page.extract_text().strip()),
                "text_length": len(page.extract_text()),
            }

            # Check for embedded objects
            if "/Annots" in page:
                page_info["has_annotations"] = True
            if "/XObject" in page.get("/Resources", {}):
                page_info["has_embedded_objects"] = True

            pages_info.append(page_info)

        report["structure_analysis"] = {
            "pages_analyzed": len(pages_info),
            "pages": pages_info,
        }

        # Check for JavaScript (potential security issue)
        try:
            catalog = reader.trailer.get("/Root", {})
            if hasattr(catalog, "get_object"):
                catalog = catalog.get_object()
            if "/OpenAction" in catalog or "/AA" in catalog:
                report["anomalies"].append(
                    "PDF contains auto-execute actions (OpenAction/AA)"
                )
            if "/JavaScript" in catalog or "/JS" in catalog:
                report["anomalies"].append(
                    "PDF contains JavaScript - potential security concern"
                )
        except Exception:
            pass

        # Check for embedded files
        try:
            if "/Names" in reader.trailer.get("/Root", {}):
                report["anomalies"].append(
                    "PDF contains named objects - may include embedded files"
                )
        except Exception:
            pass

        if reader.is_encrypted:
            report["anomalies"].append("PDF is encrypted/password-protected")

    except ImportError:
        report["document_properties"]["error"] = "PyPDF2 not available"
    except Exception as e:
        report["document_properties"]["error"] = str(e)


def _analyze_docx(file_path, report):
    """DOCX-specific analysis."""
    try:
        from docx import Document

        doc = Document(file_path)

        report["document_properties"] = {
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            "section_count": len(doc.sections),
            "style_count": len(doc.styles),
        }

        # Core properties (metadata)
        props = doc.core_properties
        metadata = {}
        for attr in [
            "author", "category", "comments", "content_status",
            "created", "identifier", "keywords", "language",
            "last_modified_by", "last_printed", "modified",
            "revision", "subject", "title", "version",
        ]:
            val = getattr(props, attr, None)
            if val:
                metadata[attr] = str(val)

        report["metadata"] = metadata if metadata else {"note": "No metadata"}

        # Content analysis
        total_chars = sum(len(p.text) for p in doc.paragraphs)
        total_words = sum(len(p.text.split()) for p in doc.paragraphs)

        # Check for embedded images
        image_count = 0
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                image_count += 1

        report["structure_analysis"] = {
            "total_characters": total_chars,
            "total_words": total_words,
            "embedded_images": image_count,
            "has_headers_footers": any(
                s.header.is_linked_to_previous is False
                for s in doc.sections
            ),
        }

        if image_count > 20:
            report["anomalies"].append(
                f"High number of embedded images: {image_count}"
            )

        # Check for macros (VBA) - look in ZIP for vbaProject.bin
        try:
            with zipfile.ZipFile(file_path, "r") as zf:
                for name in zf.namelist():
                    if "vbaProject" in name:
                        report["anomalies"].append(
                            "Document contains VBA macros - potential security concern"
                        )
                        break
        except Exception:
            pass

    except ImportError:
        report["document_properties"]["error"] = "python-docx not available"
    except Exception as e:
        report["document_properties"]["error"] = str(e)


def _analyze_xlsx(file_path, report):
    """XLSX-specific analysis."""
    try:
        from openpyxl import load_workbook

        wb = load_workbook(file_path, read_only=True, data_only=True)

        sheets_info = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            sheets_info.append({
                "name": sheet_name,
                "max_row": ws.max_row,
                "max_column": ws.max_column,
            })

        report["document_properties"] = {
            "sheet_count": len(wb.sheetnames),
            "sheet_names": wb.sheetnames,
        }

        report["structure_analysis"] = {
            "sheets": sheets_info,
        }

        # Metadata
        props = wb.properties
        metadata = {}
        for attr in [
            "creator", "title", "subject", "description",
            "created", "modified", "lastModifiedBy",
            "category", "keywords", "version",
        ]:
            val = getattr(props, attr, None)
            if val:
                metadata[attr] = str(val)

        report["metadata"] = metadata if metadata else {"note": "No metadata"}

        # Check for macros
        try:
            with zipfile.ZipFile(file_path, "r") as zf:
                for name in zf.namelist():
                    if "vbaProject" in name:
                        report["anomalies"].append(
                            "Spreadsheet contains VBA macros - "
                            "potential security concern"
                        )
                        break
        except Exception:
            pass

        wb.close()

    except ImportError:
        report["document_properties"]["error"] = "openpyxl not available"
    except Exception as e:
        report["document_properties"]["error"] = str(e)


def _analyze_text(file_path, report):
    """Plain text file analysis."""
    file_size = report["file_info"]["file_size_bytes"]

    # Read up to 1MB for analysis
    read_size = min(file_size, 1024 * 1024)
    with open(file_path, "rb") as f:
        raw_data = f.read(read_size)

    # Detect encoding
    encoding = "unknown"
    bom_detected = None
    if raw_data[:3] == b"\xef\xbb\xbf":
        encoding = "UTF-8 (BOM)"
        bom_detected = "UTF-8 BOM"
    elif raw_data[:2] == b"\xff\xfe":
        encoding = "UTF-16 LE (BOM)"
        bom_detected = "UTF-16 LE BOM"
    elif raw_data[:2] == b"\xfe\xff":
        encoding = "UTF-16 BE (BOM)"
        bom_detected = "UTF-16 BE BOM"
    else:
        try:
            raw_data.decode("utf-8")
            encoding = "UTF-8 (no BOM)"
        except UnicodeDecodeError:
            try:
                raw_data.decode("latin-1")
                encoding = "Latin-1/ISO-8859-1"
            except Exception:
                encoding = "Binary/Unknown"

    # Line analysis
    try:
        text = raw_data.decode(
            "utf-8" if "UTF-8" in encoding else "latin-1",
            errors="replace"
        )
        lines = text.splitlines()

        report["document_properties"] = {
            "encoding": encoding,
            "bom": bom_detected,
            "line_count": len(lines),
            "character_count": len(text),
            "word_count": len(text.split()),
            "line_ending": _detect_line_ending(raw_data),
        }
    except Exception as e:
        report["document_properties"] = {
            "encoding": encoding,
            "error": str(e),
        }

    # Check for null bytes in text files (binary data mixed in)
    null_count = raw_data.count(b"\x00")
    if null_count > 0 and "UTF-16" not in encoding:
        report["anomalies"].append(
            f"Null bytes found in text file: {null_count} occurrences "
            f"- possible binary data embedded"
        )


def _analyze_zip_structure(file_path, report):
    """Analyze the internal ZIP structure of Office documents."""
    try:
        with zipfile.ZipFile(file_path, "r") as zf:
            entries = []
            total_compressed = 0
            total_uncompressed = 0

            for info in zf.infolist():
                entries.append({
                    "name": info.filename,
                    "compressed_size": info.compress_size,
                    "uncompressed_size": info.file_size,
                    "compression_method": (
                        "Deflated" if info.compress_type == zipfile.ZIP_DEFLATED
                        else "Stored" if info.compress_type == zipfile.ZIP_STORED
                        else f"Method_{info.compress_type}"
                    ),
                })
                total_compressed += info.compress_size
                total_uncompressed += info.file_size

            zip_info = {
                "total_entries": len(entries),
                "total_compressed": total_compressed,
                "total_uncompressed": total_uncompressed,
                "compression_ratio": (
                    round(total_compressed / total_uncompressed, 4)
                    if total_uncompressed > 0 else 0
                ),
                "entries": entries[:50],  # Cap at 50 entries for readability
            }

            if "structure_analysis" not in report:
                report["structure_analysis"] = {}
            report["structure_analysis"]["zip_structure"] = zip_info

            # Check for suspicious entries
            for entry in entries:
                if entry["name"].startswith("..") or "/" in entry["name"]:
                    if entry["name"].startswith("../"):
                        report["anomalies"].append(
                            f"Path traversal in ZIP entry: {entry['name']}"
                        )

    except zipfile.BadZipFile:
        report["anomalies"].append(
            "File has ZIP extension but is not a valid ZIP archive"
        )
    except Exception as e:
        if "structure_analysis" not in report:
            report["structure_analysis"] = {}
        report["structure_analysis"]["zip_error"] = str(e)


def _detect_line_ending(data):
    """Detect the predominant line ending style."""
    crlf = data.count(b"\r\n")
    lf = data.count(b"\n") - crlf
    cr = data.count(b"\r") - crlf

    if crlf > lf and crlf > cr:
        return "CRLF (Windows)"
    elif lf > cr:
        return "LF (Unix/Mac)"
    elif cr > 0:
        return "CR (Old Mac)"
    return "No line endings detected"
